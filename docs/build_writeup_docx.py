"""Build a polished .docx version of docs/writeup.md.

Output: docs/writeup_polished.docx
US Letter, Arial 11pt, structured headings, styled tables, page numbers.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Setup ----------
doc = Document()

section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

normal = doc.styles['Normal']
normal.font.name = 'Arial'
normal.font.size = Pt(11)

for h_id, font_size, space_before, space_after in [
    ('Heading 1', 16, 18, 6),
    ('Heading 2', 13, 12, 4),
    ('Heading 3', 11.5, 8, 2),
]:
    s = doc.styles[h_id]
    s.font.name = 'Arial'
    s.font.size = Pt(font_size)
    s.font.bold = True
    s.font.color.rgb = RGBColor(0, 0, 0)
    s.paragraph_format.space_before = Pt(space_before)
    s.paragraph_format.space_after = Pt(space_after)


# ---------- Helpers ----------
def add_runs(paragraph, parts):
    """parts: list of (text, kind) where kind in {'plain', 'bold', 'italic', 'code'}."""
    for text, kind in parts:
        run = paragraph.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(11)
        if kind == 'bold':
            run.bold = True
        elif kind == 'italic':
            run.italic = True
        elif kind == 'code':
            run.font.name = 'Courier New'
            run.font.size = Pt(10)


def add_paragraph(parts, *, style=None, indent_left=None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    if indent_left is not None:
        p.paragraph_format.left_indent = Inches(indent_left)
    add_runs(p, parts)
    return p


def add_bullet(parts):
    p = doc.add_paragraph(style='List Bullet')
    add_runs(p, parts)
    return p


def add_code_block(text):
    """Render multi-line text in monospace (e.g., the architecture diagram)."""
    for line in text.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
    # blank line after
    doc.add_paragraph()


def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def add_styled_table(headers, rows, col_widths_in=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.autofit = False

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.font.name = 'Arial'
        run.font.size = Pt(10.5)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(cell, '2E5C8A')

    # Body rows
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ''
            # support inline parts spec or plain string
            if isinstance(val, list):
                add_runs(cell.paragraphs[0], val)
            else:
                run = cell.paragraphs[0].add_run(str(val))
                run.font.name = 'Arial'
                run.font.size = Pt(10)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Column widths
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_page_number_to_footer():
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run()
    fld_begin = OxmlElement('w:fldChar')
    fld_begin.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'PAGE'
    fld_end = OxmlElement('w:fldChar')
    fld_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    run.font.name = 'Arial'
    run.font.size = Pt(9)


# ---------- Content ----------

# Title
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(2)
run = t.add_run('Macro Event Ripple Tracker — Project Report')
run.font.name = 'Arial'
run.font.size = Pt(18)
run.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.paragraph_format.space_after = Pt(18)
sub_run = sub.add_run('Applied Finance, v0.2 · Fordham University · 2026-04-28')
sub_run.font.name = 'Arial'
sub_run.font.size = Pt(11)
sub_run.italic = True
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


# 1. Business Problem
doc.add_heading('1. Business Problem', level=1)
add_paragraph([
    ("When a major macro or geopolitical event happens — a sanctions package, a strait closure, an OPEC+ surprise cut — analysts need to map the second-order ripples across industries and asset classes within hours, not days. Yet the inputs are scattered: hundreds of news articles per week, dozens of correlated tickers, and a moving timeline of sub-events that interact non-linearly. Doing this by hand is slow and error-prone; missing a ripple costs alpha or risk-management lead time.", 'plain'),
])
add_paragraph([
    ("Macro Event Ripple Tracker (MERT)", 'bold'),
    (" is an agentic RAG system that turns a single event description into a grounded ripple analysis: a chronological timeline, a multi-level industry impact tree, attached market data, and a free-form Q&A surface — all with citations back to source articles. The MVP is wired to one event (2026 Iran War / Strait of Hormuz closure), but every component is event-agnostic and driven by a YAML config, so a new crisis is a config copy plus one ingestion run away.", 'plain'),
])


# 2. Architecture
doc.add_heading('2. Architecture', level=1)

add_code_block("""┌──────────────────────────────────────────────────────────────────┐
│  Streamlit UI (M5)                                                │
│  ┌─────────────┐ ┌──────────────┐ ┌────────────────────────┐     │
│  │ Price chart │ │ Detail panel │ │ Sidebar persistent chat │    │
│  └─────────────┘ └──────────────┘ └────────────────────────┘     │
│  ┌──────────────┐ ┌──────────────────────────────────────┐       │
│  │  Event axis  │ │  Ripple tree (interactive, agraph)    │      │
│  └──────────────┘ └──────────────────────────────────────┘       │
└──────────────────────────────────────────────────────────────────┘
                                │
                                ▼
┌──────────────── Supervisor (M4, LangGraph) ─────────────────────┐
│  classify_intent → { qa | timeline | market | ripple }            │
│         │                                                          │
│  ┌──────┴──────┬───────────┬────────────────┐                     │
│  ▼             ▼           ▼                ▼                      │
│ run_qa     run_news   run_market       run_ripple                 │
│  agent      agent      agent          (M3 generator)              │
└─────────┬───────────────────┬────────────────┬───────────────────┘
          │                   │                │
          ▼                   ▼                ▼
┌──────────────────┐  ┌─────────────────┐  ┌──────────────────┐
│ Vector retrieval │  │   Market data    │  │  Ripple tree LLM │
│ (ChromaDB +      │  │   (yfinance      │  │  (Claude Sonnet) │
│  MiniLM)         │  │    CSV cache)    │  │                  │
└────────▲─────────┘  └────────▲────────┘  └──────────────────┘
         │                     │
┌────────┴─────────┐  ┌────────┴────────┐
│ M1 news ingest   │  │ M2 market data  │
│ (GDELT+NewsAPI)  │  │ (yfinance)      │
└──────────────────┘  └─────────────────┘""")

add_paragraph([
    ("Tech stack. ", 'bold'),
    ("Python 3.11; Claude Sonnet 4.6 via ", 'plain'),
    ("langchain-anthropic", 'code'),
    (" + ", 'plain'),
    ("langgraph", 'code'),
    ("; ChromaDB + sentence-transformers ", 'plain'),
    ("all-MiniLM-L6-v2", 'code'),
    (" (local, free) for dense retrieval; Streamlit + Plotly + streamlit-agraph for the UI; yfinance + pandas for prices. All news sources are free-tier (per course constraint).", 'plain'),
])

add_paragraph([("Key design decisions.", 'bold')])
add_bullet([
    ("Event config drives every event-specific knob — keywords, tickers, date windows live in ", 'plain'),
    ("events/iran_war.yaml", 'code'),
    ("; production code holds zero hardcoded event references. Adding a new event = copying the YAML.", 'plain'),
])
add_bullet([
    ("Modular per-zone UI — ", 'plain'),
    ("ui/price_chart.py", 'code'),
    (", ", 'plain'),
    ("ui/event_axis.py", 'code'),
    (", ", 'plain'),
    ("ui/ripple.py", 'code'),
    (", ", 'plain'),
    ("ui/sidebar_chat.py", 'code'),
    (", ", 'plain'),
    ("ui/price_detail_panel.py", 'code'),
    (" each export a single ", 'plain'),
    ("render(cfg, as_of)", 'code'),
    (" so layouts can be re-arranged without touching logic, and pure helpers are unit-testable in isolation.", 'plain'),
])
add_bullet([
    ("LangGraph for the supervisor loop — ", 'plain'),
    ("agent_supervisor.py", 'code'),
    (" classifies user intent, routes to one of four specialized workers, and merges partial state. Workers can be tested independently with mocked LLMs.", 'plain'),
])
add_bullet([
    ("Three-phase ripple generation — ", 'plain'),
    ("agent_ripple.py", 'code'),
    (" first asks the LLM for the bare tree structure, then attaches news citations per node via vector retrieval, then attaches market data. Avoids one giant LLM call and keeps each phase auditable.", 'plain'),
])
add_bullet([
    ("Honest fallbacks at every LLM-JSON boundary — every consumer guards both ", 'plain'),
    ("json.JSONDecodeError", 'code'),
    (" and an ", 'plain'),
    ("isinstance", 'code'),
    (" shape gate. Wrong-shape JSON is a strictly more likely failure than parse failure when the model is instructed to “emit JSON,” so both must be handled.", 'plain'),
])


# 3. Dataset
doc.add_heading('3. Dataset', level=1)
add_paragraph([
    ("News corpus — 1,387 unique articles", 'bold'),
    (" over the event window 2026-02-28 → 2026-04-16, deduplicated across three sources:", 'plain'),
])

add_styled_table(
    headers=['Source', 'Raw count', 'Mechanism'],
    rows=[
        ['GDELT 2.0 DOC API', '1,750', '7 weekly chunks × 250-record cap (GDELT hard ceiling per query)'],
        ['NewsAPI.org (free)', '100', '100-result hard cap on free tier; 30-day lookback'],
        ['RSS (Reuters/AP)', '0', 'Feeds shut down June 2020; kept as pluggable but empty'],
        ['After MinHash LSH dedup', '1,387 unique', 'URL dedup: −6; near-duplicate (Jaccard ≥ 0.95): −457'],
    ],
    col_widths_in=[2.0, 1.2, 3.2],
)

add_paragraph([
    ("Market data — 11 tickers", 'bold'),
    (" — daily OHLCV from yfinance, cached as CSV. Includes Brent (BZ=F), WTI (CL=F), natural gas (NG=F), aluminum (ALI=F), S&P 500 (^GSPC), defense ETF (ITA), energy ETF (XLE), CF Industries (fertilizer), plus three more spanning shipping / petrochemicals / utilities.", 'plain'),
])
add_paragraph([
    ("Embeddings.", 'bold'),
    (" ChromaDB persistent index over ", 'plain'),
    ("headline + snippet", 'code'),
    (" text, embedded via ", 'plain'),
    ("all-MiniLM-L6-v2", 'code'),
    (" (local sentence-transformer, 256-token input window, ~80MB on first download). Embedding model is free and runs entirely on CPU.", 'plain'),
])
add_paragraph([
    ("Why free-only.", 'bold'),
    (" The course rule excludes paid APIs (Mediastack, NewsCatcher, Bloomberg, Refinitiv) and any full-text scraping that would touch ToS gray-area domains. The dataset section's biggest single consequence is that the corpus is built over headlines plus ~200-character previews rather than full article bodies — see Limitations §5.1.", 'plain'),
])


# 4. Evaluation
doc.add_heading('4. Evaluation', level=1)
add_paragraph([
    ("Eval harness in ", 'plain'),
    ("eval/", 'code'),
    (" measures four dimensions per spec §9, run end-to-end via ", 'plain'),
    ("python -m eval.run_eval --event iran_war", 'code'),
    (". Per Lecture 9 Slide 10's Week 2/3 emphasis on iterative quality-testing, the system was iterated ", 'plain'),
    ("six times", 'bold'),
    (" (v1 → v6) over 2026-04-27/28, with each iteration committed as a code/prompt change paired with a markdown report.", 'plain'),
])

doc.add_heading('v1 → v6 numbers', level=2)
add_styled_table(
    headers=['Dimension', 'v1', 'v6', 'Δ', 'What drove the change'],
    rows=[
        ['§9.1 retrieval precision@5', '0.76', '0.76', '—', '(plateau is structural)'],
        ['§9.2 ripple sector precision', '21.2%', [('58.6%', 'bold')], [('+37.4 pp', 'bold')], 'Scoring bug fix + token-overlap matcher'],
        ['§9.2 ripple sector recall', '58.3%', [('75.0%', 'bold')], [('+16.7 pp', 'bold')], 'Token-overlap matcher + corpus refresh'],
        ['§9.2 price integrity (in tree)', '33/33', '36/36', '—', 'always perfect'],
        ['§9.3 QA faithfulness', '0.55', [('0.60', 'bold')], '+0.05', 'Corpus refresh broke out of judge noise band'],
        ['§9.4 market spot-check', '5/5', '5/5', '—', 'always perfect'],
    ],
    col_widths_in=[2.1, 0.5, 0.5, 0.7, 2.6],
)

doc.add_heading('Iteration narrative', level=2)
add_paragraph([
    ("Iteration 1 — §9.2 N:1 scoring bug.", 'bold'),
    (" The original ", 'plain'),
    ("score()", 'code'),
    (" loop iterated truth sectors first, with an early ", 'plain'),
    ("break", 'code'),
    (" after the first AI-sector match. Once one AI sector matched a given truth, every sibling AI sector that referred to the same truth was incorrectly labelled “hallucinated.” Fix: invert the loop so each AI sector is scored independently against all truths. Precision: 21.2% → 37.1%.", 'plain'),
])
add_paragraph([
    ("Iteration 2 — token-overlap fuzzy matcher.", 'bold'),
    (" Substring-only matching missed obvious paraphrases — “Industrial Metals & Materials” was hallucinated against “Aluminum / energy-intensive metals” because no literal substring matched. Added a fallback: lowercase + plural-strip + a small stoplist of generic business words (industry, manufacturing, service, market, system, supply, equipment, global, energy) + token Jaccard. Precision: 37.1% → 48.1%; recall: 58.3% → 66.7%.", 'plain'),
])
add_paragraph([
    ("Iteration 3 — QA prompt tightening.", 'bold'),
    (" Added four hardening rules to ", 'plain'),
    ("prompts/qa_system.txt", 'code'),
    (" forbidding extrapolation (no inventing numbers / prices, prefer snippet-close phrasing, offer “available coverage indicates X but does not specify Y” instead of filling gaps). Empirical impact on §9.3 alone was within LLM-judge noise (0.50–0.55 range). Kept anyway because the rules are forward-compatible with any future corpus expansion.", 'plain'),
])
add_paragraph([
    ("Iteration 4 — corpus refresh.", 'bold'),
    (" Re-fetched GDELT + NewsAPI to lift the corpus from 1,031 → 1,387 unique articles. §9.3 jumped to ", 'plain'),
    ("0.60", 'bold'),
    (" — the first real movement on §9.3, breaking out of the noise band. §9.2 recall to 75.0%. §9.1 precision@5 stayed at 0.76, which ", 'plain'),
    ("proves corpus size was not the §9.1 bottleneck", 'bold'),
    (".", 'plain'),
])
add_paragraph([
    ("Iteration 5 — query rewriting.", 'bold'),
    (" Hypothesis: §9.1 stuck at 0.76 because retrieval queries were too generic. Added ", 'plain'),
    ("eval/query_rewriter.py", 'code'),
    (" to LLM-rewrite each test query with event context before retrieval (e.g., ", 'plain'),
    ("“How high did Brent crude go after Hormuz closed?”", 'italic'),
    (" became ", 'plain'),
    ("“Brent crude oil price spike Hormuz closure Iran war 2026 shipping disruption”", 'italic'),
    ("). precision@5 stayed at 0.76. Per-query analysis showed the LLM judge consistently rejects ~24% of hits because it strictly enforces ", 'plain'),
    ("must_be_about", 'code'),
    (" keywords — articles about general “oil prices” get rejected for not naming “Brent” specifically. The plateau is ", 'plain'),
    ("structural", 'bold'),
    (", not a query-phrasing problem.", 'plain'),
])

add_paragraph([
    ("The complete v1–v6 reports are preserved at ", 'plain'),
    ("eval/results/eval-iran_war-*.{md,json}", 'code'),
    (" and form the appendix evidence for this writeup.", 'plain'),
])


# 5. Limitations
doc.add_heading('5. Limitations', level=1)

doc.add_heading('5.1 Primary: Free-tier APIs only — corpus is headline + ~200 chars, not full article text', level=2)
add_paragraph([
    ("The ingestion pipeline uses only free-tier data sources, per the course constraint. The consequence is structural:", 'plain'),
])
add_bullet([
    ("GDELT 2.0", 'bold'),
    (" DOC API returns article ", 'plain'),
    ("metadata only", 'bold'),
    (" — URL, headline, domain, publication date. No body content.", 'plain'),
])
add_bullet([
    ("NewsAPI.org free tier", 'bold'),
    (" returns a ~200-character ", 'plain'),
    ("content", 'code'),
    (" preview, capped at 100 total results per query, within a 30-day lookback.", 'plain'),
])
add_bullet([
    ("Reuters / AP RSS feeds", 'bold'),
    (" shut down in June 2020 and return zero usable content.", 'plain'),
])
add_paragraph([
    ("The vector index is therefore built primarily over ", 'plain'),
    ("headlines plus short descriptions", 'bold'),
    (", not full articles. This caps analytical depth in two specific ways: (a) the grounded QA agent cannot cite mid-article evidence, only headline-level claims — which is why §9.3 faithfulness tops out near 0.60 even with a tightened prompt; (b) the ripple tree's “supporting news” per sector is a snippet pointer rather than an excerpted quote, which is informative but not exhaustive. The system responds honestly when coverage is insufficient — the UI surfaces explicit ", 'plain'),
    ("no_nearby_news", 'code'),
    (" / ", 'plain'),
    ("insufficient_evidence", 'code'),
    (" reason codes instead of hallucinating.", 'plain'),
])
add_paragraph([
    ("Lifting this single constraint would require either (a) ", 'plain'),
    ("full-text scraping", 'bold'),
    (" — blocked by paywalls on major outlets (WSJ, FT, NYT, Bloomberg, Economist) and ToS gray-area on others, or (b) a ", 'plain'),
    ("paid news API tier", 'bold'),
    (" (Mediastack, NewsCatcher, etc.) — both excluded by the project's free-only constraint. Every other limitation in this report is downstream of this one.", 'plain'),
])

doc.add_heading('5.2 Single-event scope', level=2)
add_paragraph([
    ("All wiring is event-agnostic and YAML-driven (a new event = ", 'plain'),
    ("cp events/iran_war.yaml events/<new>.yaml", 'code'),
    (", edit keywords + tickers, run ", 'plain'),
    ("setup.py", 'code'),
    ("), but only the 2026 Iran War / Strait of Hormuz crisis is fully populated. A reference corpus from analogous historical events (1979 Iranian Revolution, 1990–91 Gulf War) was scoped as a week-2 add-on but did not ship in v0.2.", 'plain'),
])

doc.add_heading('5.3 Prompt injection surface', level=2)
add_paragraph([
    ("News snippets are interpolated directly into the QA / news worker system prompts without delimiter escaping. A hostile snippet of the form ", 'plain'),
    ("“Ignore previous instructions...”", 'italic'),
    (" would be passed verbatim to Claude. Acceptable for an MVP that ingests from reputable aggregators (worst case is a misleading citation, not exfiltration), but a production deployment would need delimiter-wrapped context blocks or an injection-pattern pre-filter.", 'plain'),
])


# 6. Next Steps
doc.add_heading('6. Next Steps', level=1)
add_paragraph([
    ("In rough priority order, all anchored to spec §11 “Future Work”:", 'plain'),
])
add_bullet([
    ("More events.", 'bold'),
    (" The architecture is already event-agnostic. The first multi-event sprint is to populate two more ", 'plain'),
    ("events/*.yaml", 'code'),
    (" files (a non-energy macro shock for breadth, a historical event for depth) and walk through the same ingestion + ripple-generation pipeline end-to-end. (§11.1)", 'plain'),
])
add_bullet([
    ("Historical reference corpus", 'bold'),
    (" — curated summaries from the ", 'plain'),
    ("1979 Iranian Revolution", 'bold'),
    (" and ", 'plain'),
    ("1990–91 Gulf War", 'bold'),
    (", loaded by ", 'plain'),
    ("agent_ripple.generate_structure", 'code'),
    (" as few-shot priors. The goal is not a full pipeline for those events but a small set of hand-edited markdown files that anchor sector mechanics in historical analogy when the LLM generates the ripple tree.", 'plain'),
])
add_bullet([
    ("Multi-event side-by-side comparison.", 'bold'),
    (" Once two events exist, a comparison dashboard lets analysts hold sector mechanics constant and see how a different event reshapes the tree (which sectors recur, which severity rankings invert).", 'plain'),
])
add_bullet([
    ("User-input arbitrary events.", 'bold'),
    (" A “New Event” UI form replacing the YAML edit step, validating ticker symbols and date ranges before kicking off ", 'plain'),
    ("setup.py", 'code'),
    (". (§11.1)", 'plain'),
])
add_bullet([
    ("Full-text article access.", 'bold'),
    (" The single biggest unblocker for analytical depth — see Limitations §5.1. Two paths: (a) a paid API tier outside the course's free-data rule, or (b) targeted full-text scraping with a ToS-safe domain whitelist. Lifting this directly raises the §9.3 faithfulness ceiling. (§11.2)", 'plain'),
])
add_bullet([
    ("Real-time / incremental data refresh.", 'bold'),
    (" Today ", 'plain'),
    ("setup.py --refresh", 'code'),
    (" rebuilds the full window from scratch. Production use needs an incremental path — daily GDELT chunks, rolling-window expiry, streaming dashboard. (§11.4)", 'plain'),
])
add_bullet([
    ("Quantitative event-study layer.", 'bold'),
    (" Per §11.3, layer Granger causality and event-study return / volatility statistics onto the qualitative ripple tree, converting “Defense / Aerospace surged” into measured Δ-returns + significance tests against a baseline.", 'plain'),
])
add_bullet([
    ("Knowledge-Graph RAG.", 'bold'),
    (" Model sector dependencies as graph edges (fertilizer ← natural gas ← oil ← geopolitical) so ripple chains are queryable rather than re-generated each time. (§11.3)", 'plain'),
])
add_bullet([
    ("Continuous eval-drift detection", 'bold'),
    (" with TruLens or equivalent, replacing the current snapshot-only harness with online quality monitoring. (§11.5)", 'plain'),
])


# Footer info block
doc.add_paragraph()  # spacer
hr = doc.add_paragraph()
hr_run = hr.add_run('—' * 50)
hr_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

add_paragraph([
    ("Repository: ", 'bold'),
    ("github.com/fangyihe-fordham/macro-ripple-tracker", 'plain'),
])
add_paragraph([
    ("Final eval report: ", 'bold'),
    ("eval/results/eval-iran_war-20260428-003229.md", 'code'),
])
add_paragraph([
    ("Test suite: ", 'bold'),
    ("112 passed + 4 RUN_LIVE-gated", 'plain'),
])


# Page numbers
add_page_number_to_footer()


# ---------- Save ----------
doc.save('docs/writeup_polished.docx')
print('Wrote docs/writeup_polished.docx')
